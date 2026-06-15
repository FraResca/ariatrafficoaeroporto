#!/usr/bin/env python3
"""Convert a small Markdown subset to .docx using python-docx."""

from __future__ import annotations

import argparse
import re
import subprocess
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Emu, Pt
except ImportError:  # pragma: no cover
    Document = None
    Pt = None
    Emu = None


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Converte un file Markdown in DOCX.")
    parser.add_argument("input", type=Path, help="File Markdown sorgente.")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=None,
        help="File DOCX di output. Default: stesso nome del file input con estensione .docx",
    )
    return parser.parse_args()


def normalize_table_line(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def is_table_separator(line: str) -> bool:
    cells = normalize_table_line(line)
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_inline_runs(paragraph, text: str) -> None:
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) >= 3:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 3:
            run = paragraph.add_run(part[1:-1])
            if run.font is not None:
                run.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def add_code_paragraph(document, text: str) -> None:
    paragraph = document.add_paragraph(style="No Spacing")
    run = paragraph.add_run(text)
    if run.font is not None:
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def set_cell_text(cell, text: str, font_size_pt: int | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    add_inline_runs(paragraph, text)
    if font_size_pt is not None:
        for run in paragraph.runs:
            if run.font is not None:
                run.font.size = Pt(font_size_pt)


def page_text_width(document) -> int:
    section = document.sections[-1]
    return section.page_width - section.left_margin - section.right_margin


def resolve_media_path(markdown_path: Path, raw_path: str) -> Path:
    candidate = Path(raw_path)
    if candidate.is_absolute():
        return candidate
    return (markdown_path.parent / candidate).resolve()


def convert_svg_to_png(svg_path: Path, temp_dir: Path) -> Path:
    output_path = temp_dir / f"{svg_path.stem}.png"
    subprocess.run(
        [
            "inkscape",
            str(svg_path),
            "--export-type=png",
            f"--export-filename={output_path}",
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return output_path


def add_markdown_image(document, markdown_path: Path, alt_text: str, raw_path: str, temp_dir: Path) -> None:
    image_path = resolve_media_path(markdown_path, raw_path)
    if not image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[Immagine non trovata: {raw_path}]")
        return

    render_path = image_path
    if image_path.suffix.lower() == ".svg":
        render_path = convert_svg_to_png(image_path, temp_dir)

    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run()
    run.add_picture(str(render_path), width=Emu(page_text_width(document)))

    if alt_text.strip():
        caption = document.add_paragraph()
        caption.alignment = 1
        cap_run = caption.add_run(alt_text.strip())
        cap_run.italic = True


def add_markdown_table(document, lines: list[str]) -> None:
    rows = [normalize_table_line(line) for line in lines]
    if len(rows) < 2 or not is_table_separator(lines[1]):
        for raw in lines:
            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, raw)
        return

    header = rows[0]
    body = rows[2:]
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = False
    font_size = 8 if len(header) >= 8 else 9 if len(header) >= 6 else 10
    col_width = int(page_text_width(document) / max(1, len(header)))
    for idx, value in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.width = col_width
        set_cell_text(cell, value, font_size_pt=font_size)

    for row in body:
        cells = table.add_row().cells
        for idx in range(len(header)):
            cells[idx].width = col_width
            set_cell_text(cells[idx], row[idx] if idx < len(row) else "", font_size_pt=font_size)


def flush_paragraph_buffer(document, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(line.strip() for line in buffer).strip()
    if text:
        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, text)
    buffer.clear()


def convert_markdown(input_path: Path, output_path: Path) -> None:
    if Document is None:
        raise RuntimeError(
            "python-docx non e' installato. Installa la dipendenza `python-docx` e rilancia."
        )

    document = Document()
    lines = input_path.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []
    in_code_block = False
    code_buffer: list[str] = []
    in_table = False
    table_buffer: list[str] = []

    with tempfile.TemporaryDirectory(prefix="md_docx_") as temp_dir_name:
        temp_dir = Path(temp_dir_name)

        def flush_table() -> None:
            nonlocal table_buffer, in_table
            if table_buffer:
                add_markdown_table(document, table_buffer)
                table_buffer = []
            in_table = False

        for line in lines:
            stripped = line.rstrip()

            if stripped.startswith("```"):
                flush_paragraph_buffer(document, paragraph_buffer)
                flush_table()
                if in_code_block:
                    for code_line in code_buffer:
                        add_code_paragraph(document, code_line)
                    code_buffer = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_buffer.append(stripped)
                continue

            image_match = IMAGE_RE.match(stripped.strip())
            if image_match:
                flush_paragraph_buffer(document, paragraph_buffer)
                flush_table()
                add_markdown_image(
                    document,
                    input_path,
                    image_match.group(1),
                    image_match.group(2),
                    temp_dir,
                )
                continue

            if "|" in stripped and stripped.strip():
                flush_paragraph_buffer(document, paragraph_buffer)
                table_buffer.append(stripped)
                in_table = True
                continue
            if in_table:
                flush_table()

            if not stripped.strip():
                flush_paragraph_buffer(document, paragraph_buffer)
                continue

            heading_match = HEADING_RE.match(stripped)
            if heading_match:
                flush_paragraph_buffer(document, paragraph_buffer)
                level = len(heading_match.group(1))
                heading = document.add_heading(level=level)
                add_inline_runs(heading, heading_match.group(2).strip())
                continue

            if stripped.startswith("> "):
                flush_paragraph_buffer(document, paragraph_buffer)
                paragraph = document.add_paragraph(style="Intense Quote")
                add_inline_runs(paragraph, stripped[2:].strip())
                continue

            if stripped.startswith("- "):
                flush_paragraph_buffer(document, paragraph_buffer)
                paragraph = document.add_paragraph()
                paragraph.add_run("- ")
                add_inline_runs(paragraph, stripped[2:].strip())
                continue

            ordered_match = ORDERED_RE.match(stripped)
            if ordered_match:
                flush_paragraph_buffer(document, paragraph_buffer)
                paragraph = document.add_paragraph(style="List Number")
                add_inline_runs(paragraph, ordered_match.group(1).strip())
                continue

            paragraph_buffer.append(stripped)

        flush_paragraph_buffer(document, paragraph_buffer)
        flush_table()
        if code_buffer:
            for code_line in code_buffer:
                add_code_paragraph(document, code_line)
    document.save(output_path)


def main() -> int:
    args = parse_args()
    input_path = args.input
    output_path = args.output or input_path.with_suffix(".docx")
    if not input_path.exists():
        raise FileNotFoundError(f"File non trovato: {input_path}")
    convert_markdown(input_path, output_path)
    print(f"File scritto: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
